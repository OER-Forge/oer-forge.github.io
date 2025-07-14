"""
scan.py: Asset database logic for pages and files only.
"""
import os
import sqlite3

def batch_read_files(file_paths):
    """
    Reads multiple files and returns their contents as a dict: {path: content}
    Supports markdown (.md), notebook (.ipynb), docx (.docx), and other file types.
    """
    contents = {}
    for path in file_paths:
        ext = os.path.splitext(path)[1].lower()
        try:
            if ext == '.md':
                contents[path] = read_markdown_file(path)
            elif ext == '.ipynb':
                contents[path] = read_notebook_file(path)
            elif ext == '.docx':
                contents[path] = read_docx_file(path)
            else:
                contents[path] = None
        except Exception as e:
            print(f"[ERROR] Could not read {path}: {e}")
            contents[path] = None
    return contents

def read_markdown_file(path):
    """
    Reads a markdown (.md) file and returns its content as a string.
    """
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        print(f"[ERROR] Could not read markdown file {path}: {e}")
        return None

def read_notebook_file(path):
    """
    Reads a Jupyter notebook (.ipynb) file and returns its content as a dict.
    """
    import json
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"[ERROR] Could not read notebook file {path}: {e}")
        return None

def read_docx_file(path):
    """
    Reads a docx file and returns its text content as a string.
    Requires python-docx to be installed.
    """
    try:
        from docx import Document
        doc = Document(path)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        return '\n'.join(text)
    except ImportError:
        print("[ERROR] python-docx is not installed. Run 'pip install python-docx' in your environment.")
        return None
    except Exception as e:
        print(f"[ERROR] Could not read docx file {path}: {e}")
        return None

def batch_extract_assets(contents_dict, content_type, **kwargs):
    """
    Extracts assets from multiple file contents in one pass.
    contents_dict: {path: content}
    content_type: 'markdown', 'notebook', 'docx', etc.
    Returns a dict: {path: [asset_records]}
    """
    from oerforge.db_utils import insert_file_records, link_files_to_pages, get_db_connection
    assets = {}
    # Helper: MIME type mapping (media, document, and data types)
    mime_map = {
        # Images
        '.png': 'image/png',
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.gif': 'image/gif',
        '.svg': 'image/svg+xml',
        '.bmp': 'image/bmp',
        '.webp': 'image/webp',
        # Video
        '.mp4': 'video/mp4',
        '.webm': 'video/webm',
        '.mov': 'video/quicktime',
        '.avi': 'video/x-msvideo',
        '.mkv': 'video/x-matroska',
        # Audio
        '.mp3': 'audio/mpeg',
        '.wav': 'audio/wav',
        '.ogg': 'audio/ogg',
        '.flac': 'audio/flac',
        # Documents
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
        # Data files
        '.csv': 'text/csv',
        '.tsv': 'text/tab-separated-values',
        '.json': 'application/json',
        '.xml': 'application/xml',
        '.npy': 'application/octet-stream',
        '.txt': 'text/plain',
        '.zip': 'application/zip',
        '.tar': 'application/x-tar',
        '.gz': 'application/gzip',
        '.rst': 'text/x-rst',
        # Markdown/Notebook
        '.md': 'text/markdown',
        '.ipynb': 'application/x-ipynb+json',
    }
    # Insert each source file as a page if not present
    conn = get_db_connection()
    cursor = conn.cursor()
    # Add mime_type column to content if not present
    cursor.execute("PRAGMA table_info(content)")
    columns = [row[1] for row in cursor.fetchall()]
    if 'mime_type' not in columns:
        try:
            cursor.execute("ALTER TABLE content ADD COLUMN mime_type TEXT")
        except Exception:
            pass
    for source_path in contents_dict:
        ext = os.path.splitext(source_path)[1].lower()
        mime_type = mime_map.get(ext, '')
        cursor.execute("SELECT id FROM content WHERE source_path=?", (source_path,))
        if cursor.fetchone() is None:
            cursor.execute("INSERT INTO content (source_path, output_path, is_autobuilt, mime_type) VALUES (?, ?, ?, ?)", (source_path, None, 0, mime_type))
    conn.commit()
    # Extract assets for each file type
    for path, content in contents_dict.items():
        ext = os.path.splitext(path)[1].lower()
        if content_type == 'markdown':
            assets[path] = [a for a in extract_linked_files_from_markdown_content(content, page_id=None)
                            if mime_map.get(os.path.splitext(a.get('path',''))[1].lower())]
        elif content_type == 'notebook':
            if content and isinstance(content, dict) and 'cells' in content:
                cell_assets = []
                for cell in content['cells']:
                    cell_assets.extend([a for a in extract_linked_files_from_notebook_cell_content(cell, nb_path=path)
                                       if mime_map.get(os.path.splitext(a.get('path',''))[1].lower())])
                assets[path] = cell_assets
            else:
                assets[path] = []
        elif content_type == 'docx':
            assets[path] = [a for a in extract_linked_files_from_docx_content(path, page_id=None)
                            if mime_map.get(os.path.splitext(a.get('path',''))[1].lower())] if content else []
        else:
            assets[path] = []
    # Insert asset records into files table
    file_records = []
    file_page_links = []
    for source_path, asset_list in assets.items():
        for asset in asset_list:
            asset_path = asset.get('path', '')
            asset_ext = os.path.splitext(asset_path)[1].lower()
            mime_type = mime_map.get(asset_ext, '')
            file_record = {
                'filename': os.path.basename(asset_path),
                'extension': asset_ext,
                'mime_type': mime_type,
                'is_image': int(asset_ext in ['.png','.jpg','.jpeg','.gif','.svg']),
                'is_remote': int(asset_path.startswith('http')),
                'url': asset_path,
                'referenced_page': source_path,
                'relative_path': asset_path,
                'absolute_path': None,
                'cell_type': asset.get('type', None),
                'is_code_generated': None,
                'is_embedded': None
            }
            file_records.append(file_record)
    file_ids = insert_file_records(file_records)
    # Link files to pages
    idx = 0
    for source_path, asset_list in assets.items():
        for _ in asset_list:
            file_page_links.append((file_ids[idx], source_path))
            idx += 1
    if file_page_links:
        link_files_to_pages(file_page_links)
    conn.close()
    return assets

def extract_linked_files_from_markdown_content(md_text, page_id=None):
    """
    Extracts asset links from markdown text.
    Returns a list of file records.
    """
    import re
    asset_pattern = re.compile(r'!\[[^\]]*\]\(([^)]+)\)|\[[^\]]*\]\(([^)]+)\)')
    assets = []
    for match in asset_pattern.finditer(md_text):
        asset_path = match.group(1) or match.group(2)
        if asset_path:
            assets.append({
                'type': 'asset',
                'path': asset_path,
                'page_id': page_id
            })
    return assets

def extract_linked_files_from_notebook_cell_content(cell, nb_path=None):
    """
    Extracts asset links from a notebook cell.
    Returns a list of file records.
    """
    assets = []
    # Extract markdown-linked images
    if cell.get('cell_type') == 'markdown':
        source = cell.get('source', [])
        if isinstance(source, list):
            text = ''.join(source)
        else:
            text = str(source)
        assets.extend(extract_linked_files_from_markdown_content(text, page_id=None))
        for asset in assets:
            asset['notebook'] = nb_path
    # Extract embedded/code-produced images from outputs
    if cell.get('cell_type') == 'code' and 'outputs' in cell:
        for idx, output in enumerate(cell['outputs']):
            # Typical image output: {'data': {'image/png': ...}, ...}
            if 'data' in output:
                for img_type in ['image/png', 'image/jpeg', 'image/gif', 'image/svg+xml']:
                    if img_type in output['data']:
                        ext = {
                            'image/png': '.png',
                            'image/jpeg': '.jpg',
                            'image/gif': '.gif',
                            'image/svg+xml': '.svg',
                        }[img_type]
                        nb_name = os.path.basename(nb_path) if nb_path else 'notebook'
                        rel_path = f'notebook_embedded/{nb_name}/cell{idx}{ext}'
                        assets.append({
                            'type': 'asset',
                            'path': rel_path,
                            'notebook': nb_path,
                            'filename': f'cell{idx}{ext}',
                            'extension': ext,
                            'is_embedded': True,
                            'is_code_generated': True
                        })
    return assets

def extract_linked_files_from_docx_content(docx_path, page_id=None):
    """
    Extracts asset links from a docx file.
    Returns a list of file records.
    """
    assets = []
    try:
        from docx import Document
        doc = Document(docx_path)
        import re
        asset_pattern = re.compile(r'(https?://[^\s]+|assets/[^\s]+|images/[^\s]+)')
        # Extract text-based links as before
        for para in doc.paragraphs:
            matches = asset_pattern.findall(para.text)
            for asset_path in matches:
                assets.append({
                    'type': 'asset',
                    'path': asset_path,
                    'page_id': page_id
                })
        # Extract embedded images
        for rel in doc.part.rels.values():
            if rel.reltype == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image':
                img_part = rel.target_part
                img_name = os.path.basename(img_part.partname)
                img_ext = os.path.splitext(img_name)[1].lower()
                # Use a relative path for DB, e.g. 'docx_embedded/<docx_filename>/<img_name>'
                rel_path = f'docx_embedded/{os.path.basename(docx_path)}/{img_name}'
                assets.append({
                    'type': 'asset',
                    'path': rel_path,
                    'page_id': page_id,
                    'filename': img_name,
                    'extension': img_ext,
                    'is_embedded': True
                })
    except Exception as e:
        print(f"[ERROR] Could not extract assets from docx {docx_path}: {e}")
    return assets

def populate_site_info_from_config(config_path):
    """
    Reads _config.yml and populates the site_info table with site and footer info.
    """
    import yaml
    project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    db_path = os.path.join(project_root, 'db', 'sqlite.db')
    full_config_path = os.path.join(project_root, config_path)
    with open(full_config_path, 'r', encoding='utf-8') as f:
        config = yaml.safe_load(f)
    site = config.get('site', {})
    footer = config.get('footer', {})
    theme = site.get('theme', {})
    # Read header.html contents
    header_html_path = os.path.join(project_root, 'static', 'templates', 'header.html')
    try:
        with open(header_html_path, 'r', encoding='utf-8') as hf:
            header_html = hf.read()
    except Exception:
        header_html = ''
    print('[DEBUG] header_html read from file (first 500 chars):', repr(header_html)[:500])
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM site_info")
    cursor.execute(
        """
        INSERT INTO site_info (title, author, description, logo, favicon, theme_default, theme_light, theme_dark, language, github_url, footer_text, header)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            site.get('title', ''),
            site.get('author', ''),
            site.get('description', ''),
            site.get('logo', ''),
            site.get('favicon', ''),
            theme.get('default', ''),
            theme.get('light', ''),
            theme.get('dark', ''),
            site.get('language', ''),
            site.get('github_url', ''),
            footer.get('text', ''),
            header_html
        )
    )
    conn.commit()
    conn.close()

def parse_toc_and_populate_pages(config_path):
    """
    Parses the toc: from _config.yml and populates the pages table with hierarchy, menu, autogen, parent, order, and depth info.
    Logs missing and duplicate files.
    """
    import yaml
    project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    db_path = os.path.join(project_root, 'db', 'sqlite.db')
    full_config_path = os.path.join(project_root, config_path)
    with open(full_config_path, 'r', encoding='utf-8') as f:
        config = yaml.safe_load(f)
    toc = config.get('toc', [])
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM pages")
    seen_paths = set()

    def walk_toc(items, parent_id=None, depth=0):
        for idx, item in enumerate(items):
            title = item.get('title', '')
            file_path = item.get('file')
            is_menu_item = int(item.get('menu', False))
            is_autogen = int(file_path is None)
            was_autogenerated = 0  # Set to 1 if/when actually generated
            order = idx
            # Insert page record
            if file_path:
                source_path = file_path
                output_path = None
                if source_path in seen_paths:
                    log_event(f"TOC: Duplicate file path '{source_path}' in toc", level="WARN")
                seen_paths.add(source_path)
                # Check if file exists
                abs_path = os.path.join(project_root, source_path)
                if not os.path.exists(abs_path):
                    log_event(f"TOC: Missing file '{source_path}' (expected at {abs_path})", level="ERROR")
                cursor.execute(
                    """
                    INSERT INTO pages (title, source_path, output_path, is_autobuilt, is_menu_item, is_autogen, was_autogenerated, parent_id, "order", depth)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (title, source_path, output_path, 0, is_menu_item, is_autogen, was_autogenerated, parent_id, order, depth)
                )
                page_id = cursor.lastrowid
            else:
                # Autogenerated node
                cursor.execute(
                    """
                    INSERT INTO pages (title, source_path, output_path, is_autobuilt, is_menu_item, is_autogen, was_autogenerated, parent_id, "order", depth)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (title, None, None, 0, is_menu_item, is_autogen, was_autogenerated, parent_id, order, depth)
                )
                page_id = cursor.lastrowid
            # Recursively process children
            children = item.get('children', [])
            if children:
                walk_toc(children, parent_id=page_id, depth=depth+1)

    walk_toc(toc, parent_id=None, depth=0)
    conn.commit()
    conn.close()